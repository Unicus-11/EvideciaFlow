"""
File Processor for Research Platform
Handles file uploads, validation, and processing for papers and figures
"""

import os
import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime
import logging

# Image processing
try:
    from PIL import Image, ImageStat
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Document processing
try:
    import docx
    import mammoth
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FileProcessor:
    """Handles file uploads and processing for research platform"""
    
    def __init__(self, upload_path: str = "uploads/", max_file_size: int = 50 * 1024 * 1024):
        self.upload_path = Path(upload_path)
        self.upload_path.mkdir(exist_ok=True)
        self.max_file_size = max_file_size  # 50MB default
        
        # Create subdirectories
        (self.upload_path / "papers").mkdir(exist_ok=True)
        (self.upload_path / "figures").mkdir(exist_ok=True)
        (self.upload_path / "protocols").mkdir(exist_ok=True)
        (self.upload_path / "temp").mkdir(exist_ok=True)
        
        self.logger = logging.getLogger(__name__)
        
        # Allowed file types
        self.allowed_paper_types = {'.pdf', '.docx', '.doc', '.txt'}
        self.allowed_figure_types = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.pdf', '.eps', '.svg'}
        self.allowed_protocol_types = {'.pdf', '.docx', '.doc', '.txt', '.md'}
    
    def validate_file(self, file_data: bytes, filename: str, file_type: str = 'paper') -> Tuple[bool, List[str]]:
        """Validate uploaded file"""
        errors = []
        
        # Check file size
        if len(file_data) > self.max_file_size:
            errors.append(f"File too large: {len(file_data) / (1024*1024):.1f}MB (max: {self.max_file_size / (1024*1024):.1f}MB)")
        
        if len(file_data) == 0:
            errors.append("Empty file")
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        
        if file_type == 'paper' and file_ext not in self.allowed_paper_types:
            errors.append(f"Invalid file type for papers: {file_ext}. Allowed: {', '.join(self.allowed_paper_types)}")
        elif file_type == 'figure' and file_ext not in self.allowed_figure_types:
            errors.append(f"Invalid file type for figures: {file_ext}. Allowed: {', '.join(self.allowed_figure_types)}")
        elif file_type == 'protocol' and file_ext not in self.allowed_protocol_types:
            errors.append(f"Invalid file type for protocols: {file_ext}. Allowed: {', '.join(self.allowed_protocol_types)}")
        
        # Basic content validation
        if file_ext in {'.pdf'} and not file_data.startswith(b'%PDF'):
            errors.append("Invalid PDF file format")
        elif file_ext in {'.png'} and not file_data.startswith(b'\x89PNG'):
            errors.append("Invalid PNG file format")
        elif file_ext in {'.jpg', '.jpeg'} and not file_data.startswith(b'\xff\xd8\xff'):
            errors.append("Invalid JPEG file format")
        
        return len(errors) == 0, errors
    
    def save_uploaded_file(self, file_data: bytes, filename: str, session_id: str, 
                          file_type: str = 'paper') -> Tuple[bool, str, Dict[str, Any]]:
        """Save uploaded file and return file info"""
        
        # Validate file first
        is_valid, errors = self.validate_file(file_data, filename, file_type)
        if not is_valid:
            return False, f"Validation failed: {', '.join(errors)}", {}
        
        # Generate unique filename
        file_ext = Path(filename).suffix.lower()
        file_hash = hashlib.md5(file_data).hexdigest()[:12]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{session_id[:8]}_{timestamp}_{file_hash}{file_ext}"
        
        # Determine subdirectory
        subdir = file_type + 's'  # papers, figures, protocols
        file_path = self.upload_path / subdir / unique_filename
        
        try:
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Get file info
            file_info = {
                'original_filename': filename,
                'saved_filename': unique_filename,
                'file_path': str(file_path),
                'file_size': len(file_data),
                'file_type': file_type,
                'file_extension': file_ext,
                'upload_timestamp': datetime.now().isoformat(),
                'file_hash': file_hash
            }
            
            self.logger.info(f"File saved: {filename} -> {unique_filename}")
            return True, str(file_path), file_info
            
        except Exception as e:
            self.logger.error(f"Error saving file {filename}: {e}")
            return False, f"Error saving file: {str(e)}", {}
    
    def extract_text_from_file(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text content from various file types"""
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in {'.docx', '.doc'}:
                return self._extract_text_from_docx(file_path)
            else:
                return False, f"Unsupported file type for text extraction: {file_ext}", {}
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return False, f"Error extracting text: {str(e)}", {}
    
    def _extract_text_from_txt(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'word_count': len(text.split()),
                'char_count': len(text),
                'line_count': len(text.splitlines())
            }
            
            return True, text, metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    metadata = {
                        'word_count': len(text.split()),
                        'char_count': len(text),
                        'line_count': len(text.splitlines()),
                        'encoding_used': encoding
                    }
                    
                    return True, text, metadata
                except:
                    continue
            
            return False, "Could not decode text file with any common encoding", {}
    
    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return False, "PDF processing libraries not available", {}
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = '\n\n'.join(text_parts)
                
                metadata = {
                    'page_count': len(pdf.pages),
                    'word_count': len(full_text.split()),
                    'char_count': len(full_text),
                    'extraction_method': 'pdfplumber'
                }
                
                return True, full_text, metadata
                
        except Exception as e1:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_parts = []
                    
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
                    
                    full_text = '\n\n'.join(text_parts)
                    
                    metadata = {
                        'page_count': len(pdf_reader.pages),
                        'word_count': len(full_text.split()),
                        'char_count': len(full_text),
                        'extraction_method': 'PyPDF2'
                    }
                    
                    return True, full_text, metadata
                    
            except Exception as e2:
                return False, f"PDF extraction failed: {str(e1)}, {str(e2)}", {}
    
    def _extract_text_from_docx(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return False, "DOCX processing libraries not available", {}
        
        try:
            # Try python-docx first
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = '\n\n'.join(text_parts)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'extraction_method': 'python-docx'
            }
            
            return True, full_text, metadata
            
        except Exception as e1:
            # Fallback to mammoth
            try:
                with open(file_path, 'rb') as f:
                    result = mammoth.extract_raw_text(f)
                    full_text = result.value
                    
                    metadata = {
                        'word_count': len(full_text.split()),
                        'char_count': len(full_text),
                        'extraction_method': 'mammoth',
                        'conversion_messages': [str(msg) for msg in result.messages]
                    }
                    
                    return True, full_text, metadata
                    
            except Exception as e2:
                return False, f"DOCX extraction failed: {str(e1)}, {str(e2)}", {}
    
    def analyze_figure(self, file_path: str) -> Tuple[bool, Dict[str, Any]]:
        """Analyze figure specifications and quality"""
        
        if not PIL_AVAILABLE:
            return False, {"error": "Image processing libraries not available"}
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._analyze_pdf_figure(file_path)
            elif file_ext in {'.png', '.jpg', '.jpeg', '.tiff', '.tif'}:
                return self._analyze_image_figure(file_path)
            else:
                return False, {"error": f"Unsupported figure type: {file_ext}"}
                
        except Exception as e:
            self.logger.error(f"Error analyzing figure {file_path}: {e}")
            return False, {"error": f"Analysis failed: {str(e)}"}
    
    def _analyze_image_figure(self, file_path: Path) -> Tuple[bool, Dict[str, Any]]:
        """Analyze image figure specifications"""
        
        with Image.open(file_path) as img:
            # Basic properties
            width, height = img.size
            mode = img.mode
            format_name = img.format
            
            # Calculate DPI
            dpi = img.info.get('dpi', (72, 72))  # Default to 72 DPI if not specified
            if isinstance(dpi, tuple):
                dpi_x, dpi_y = dpi
            else:
                dpi_x = dpi_y = dpi
            
            # File size
            file_size = file_path.stat().st_size
            
            # Color analysis
            if img.mode == 'RGB':
                # Check if image is effectively grayscale
                stat = ImageStat.Stat(img)
                is_grayscale = all(abs(stat.mean[i] - stat.mean[0]) < 10 for i in range(3))
            else:
                is_grayscale = img.mode in ['L', 'P', '1']
            
            # Quality assessment
            quality_issues = []
            
            if min(dpi_x, dpi_y) < 300:
                quality_issues.append(f"Low DPI: {min(dpi_x, dpi_y)} (recommended: 300+)")
            
            if width * height > 10000 * 10000:  # Very large images
                quality_issues.append(f"Very large dimensions: {width}x{height}")
            elif width * height < 600 * 400:  # Very small images
                quality_issues.append(f"Small dimensions: {width}x{height}")
            
            if file_size > 50 * 1024 * 1024:  # 50MB
                quality_issues.append(f"Large file size: {file_size / (1024*1024):.1f}MB")
            
            analysis = {
                'width': width,
                'height': height,
                'dpi_x': dpi_x,
                'dpi_y': dpi_y,
                'min_dpi': min(dpi_x, dpi_y),
                'color_mode': mode,
                'format': format_name,
                'file_size_mb': file_size / (1024 * 1024),
                'is_grayscale': is_grayscale,
                'aspect_ratio': round(width / height, 2),
                'megapixels': round((width * height) / 1000000, 1),
                'quality_issues': quality_issues,
                'meets_300_dpi': min(dpi_x, dpi_y) >= 300,
                'suitable_for_print': min(dpi_x, dpi_y) >= 300 and file_size < 50 * 1024 * 1024
            }
            
            return True, analysis
    
    def _analyze_pdf_figure(self, file_path: Path) -> Tuple[bool, Dict[str, Any]]:
        """Analyze PDF figure specifications"""
    
    if not PDF_AVAILABLE:
        return False, {"error": "PDF processing not available"}
    
    try:
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                return False, {"error": "Empty PDF"}
            
            page = pdf.pages[0]  # Analyze first page
            
            # Get page dimensions (points: 72 points = 1 inch)
            width_pt = float(page.width)
            height_pt = float(page.height)
            
            # Convert to inches and estimate pixels at 300 DPI
            width_in = width_pt / 72
            height_in = height_pt / 72
            width_px_300dpi = int(width_in * 300)
            height_px_300dpi = int(height_in * 300)
            
            file_size = file_path.stat().st_size
            
            quality_issues = []
            
            if width_in < 3 or height_in < 3:
                quality_issues.append(f"Small figure size: {width_in:.1f}x{height_in:.1f} inches")
            
            if width_in > 20 or height_in > 20:
                quality_issues.append(f"Very large figure size: {width_in:.1f}x{height_in:.1f} inches")
            
            if file_size > 50 * 1024 * 1024:
                quality_issues.append(f"Large file size: {file_size / (1024*1024):.1f}MB")
            
            analysis = {
                'width_in': round(width_in, 2),
                'height_in': round(height_in, 2),
                'width_px_300dpi': width_px_300dpi,
                'height_px_300dpi': height_px_300dpi,
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'num_pages': len(pdf.pages),
                'has_images': bool(page.images),
                'quality_issues': quality_issues,
                'suitable_for_print': len(quality_issues) == 0
            }
            
            return True, analysis

    except Exception as e:
        self.logger.error(f"Error analyzing PDF figure {file_path}: {e}")
        return False, {"error": f"PDF figure analysis failed: {str(e)}"}
