from typing import Tuple, List, Dict, Any
from pathlib import Path
import zipfile
import hashlib
import logging
import re
import mimetypes
import shutil
import os
from datetime import datetime

# Image processing
try:
    from PIL import Image, ImageStat
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Document processing
try:
    import docx
    import mammoth
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FileProcessor:
    """Handles file uploads and processing for research platform"""
    
    def __init__(self, upload_path: str = "uploads/", max_file_size: int = 50 * 1024 * 1024):
        self.upload_path = Path(upload_path)
        self.upload_path.mkdir(exist_ok=True)
        self.max_file_size = max_file_size  # 50MB default
        
        # Create subdirectories
        (self.upload_path / "papers").mkdir(exist_ok=True)
        (self.upload_path / "figures").mkdir(exist_ok=True)
        (self.upload_path / "protocols").mkdir(exist_ok=True)
        (self.upload_path / "temp").mkdir(exist_ok=True)
        (self.upload_path / "archived").mkdir(exist_ok=True)
        (self.upload_path / "thumbnails").mkdir(exist_ok=True)
        
        self.logger = logging.getLogger(__name__)
        
        # Allowed file types
        self.allowed_paper_types = {'.pdf', '.docx', '.doc', '.txt'}
        self.allowed_figure_types = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.pdf', '.eps', '.svg'}
        self.allowed_protocol_types = {'.pdf', '.docx', '.doc', '.txt', '.md'}
        
        # File size limits by type
        self.size_limits = {
            'paper': 50 * 1024 * 1024,    # 50MB
            'figure': 100 * 1024 * 1024,   # 100MB
            'protocol': 25 * 1024 * 1024   # 25MB
        }
    
    def clean_filename(self, filename: str) -> str:
        """Clean filename for safe storage"""
        if not filename:
            return "unnamed_file"
        
        # Remove or replace problematic characters
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        filename = re.sub(r'[^\w\-_\. ]', '', filename)
        filename = re.sub(r' +', '_', filename)
        filename = filename.strip('._-')
        
        # Ensure filename isn't too long
        name, ext = os.path.splitext(filename)
        if len(name) > 100:
            name = name[:100]
        
        return name + ext
    
    def validate_file(self, file_data: bytes, filename: str, file_type: str = 'paper') -> Tuple[bool, List[str]]:
        """Validate uploaded file"""
        errors = []
        
        # Check file size
        size_limit = self.size_limits.get(file_type, self.max_file_size)
        if len(file_data) > size_limit:
            errors.append(f"File too large: {len(file_data) / (1024*1024):.1f}MB (max: {size_limit / (1024*1024):.1f}MB)")
        
        if len(file_data) == 0:
            errors.append("Empty file")
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        
        if file_type == 'paper' and file_ext not in self.allowed_paper_types:
            errors.append(f"Invalid file type for papers: {file_ext}. Allowed: {', '.join(self.allowed_paper_types)}")
        elif file_type == 'figure' and file_ext not in self.allowed_figure_types:
            errors.append(f"Invalid file type for figures: {file_ext}. Allowed: {', '.join(self.allowed_figure_types)}")
        elif file_type == 'protocol' and file_ext not in self.allowed_protocol_types:
            errors.append(f"Invalid file type for protocols: {file_ext}. Allowed: {', '.join(self.allowed_protocol_types)}")
        
        # Basic content validation
        if file_ext == '.pdf' and not file_data.startswith(b'%PDF'):
            errors.append("Invalid PDF file format")
        elif file_ext == '.png' and not file_data.startswith(b'\x89PNG'):
            errors.append("Invalid PNG file format")
        elif file_ext in {'.jpg', '.jpeg'} and not file_data.startswith(b'\xff\xd8\xff'):
            errors.append("Invalid JPEG file format")
        elif file_ext == '.docx' and not file_data.startswith(b'PK'):
            errors.append("Invalid DOCX file format")
        
        # Check for potentially malicious files
        dangerous_patterns = [
            b'<script',
            b'javascript:',
            b'vbscript:',
            b'<?php',
            b'<%'
        ]
        
        for pattern in dangerous_patterns:
            if pattern in file_data[:1024]:  # Check first 1KB
                errors.append("File contains potentially dangerous content")
                break
        
        return len(errors) == 0, errors
    
    def save_uploaded_file(self, file_data: bytes, filename: str, session_id: str, 
                          file_type: str = 'paper') -> Tuple[bool, str, Dict[str, Any]]:
        """Save uploaded file and return file info"""
        
        # Validate file first
        is_valid, errors = self.validate_file(file_data, filename, file_type)
        if not is_valid:
            return False, f"Validation failed: {', '.join(errors)}", {}
        
        # Clean filename
        cleaned_filename = self.clean_filename(filename)
        
        # Generate unique filename
        file_ext = Path(cleaned_filename).suffix.lower()
        file_hash = hashlib.md5(file_data).hexdigest()[:12]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{session_id[:8]}_{timestamp}_{file_hash}{file_ext}"
        
        # Determine subdirectory
        subdir = file_type + 's'  # papers, figures, protocols
        file_path = self.upload_path / subdir / unique_filename
        
        try:
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Generate file info
            file_info = {
                'original_filename': filename,
                'cleaned_filename': cleaned_filename,
                'saved_filename': unique_filename,
                'file_path': str(file_path),
                'relative_path': f"{subdir}/{unique_filename}",
                'file_size': len(file_data),
                'file_type': file_type,
                'file_extension': file_ext,
                'upload_timestamp': datetime.now().isoformat(),
                'file_hash': file_hash,
                'session_id': session_id
            }
            
            # Add file-specific metadata
            if file_type == 'figure':
                analysis_success, analysis = self.analyze_figure(str(file_path))
                if analysis_success:
                    file_info['figure_analysis'] = analysis
            
            self.logger.info(f"File saved: {filename} -> {unique_filename}")
            return True, str(file_path), file_info
            
        except Exception as e:
            self.logger.error(f"Error saving file {filename}: {e}")
            return False, f"Error saving file: {str(e)}", {}
    
    def extract_text_from_file(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text content from various file types"""
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in {'.docx', '.doc'}:
                return self._extract_text_from_docx(file_path)
            else:
                return False, f"Unsupported file type for text extraction: {file_ext}", {}
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return False, f"Error extracting text: {str(e)}", {}
    
    def _extract_text_from_txt(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'word_count': len(text.split()),
                'char_count': len(text),
                'line_count': len(text.splitlines()),
                'encoding': 'utf-8'
            }
            
            return True, text, metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    metadata = {
                        'word_count': len(text.split()),
                        'char_count': len(text),
                        'line_count': len(text.splitlines()),
                        'encoding': encoding
                    }
                    
                    return True, text, metadata
                except:
                    continue
            
            return False, "Could not decode text file with any common encoding", {}
    
    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return False, "PDF processing libraries not available", {}
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = '\n\n'.join(text_parts)
                
                if full_text.strip():
                    metadata = {
                        'page_count': len(pdf.pages),
                        'word_count': len(full_text.split()),
                        'char_count': len(full_text),
                        'extraction_method': 'pdfplumber'
                    }
                    
                    return True, full_text, metadata
                
        except Exception as e1:
            self.logger.warning(f"pdfplumber failed: {e1}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_parts = []
                
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                
                full_text = '\n\n'.join(text_parts)
                
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'word_count': len(full_text.split()),
                    'char_count': len(full_text),
                    'extraction_method': 'PyPDF2'
                }
                
                return True, full_text, metadata
                
        except Exception as e2:
            return False, f"PDF extraction failed: {str(e2)}", {}
    
    def _extract_text_from_docx(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return False, "DOCX processing libraries not available", {}
        
        try:
            # Try python-docx first
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = '\n\n'.join(text_parts)
            
            if full_text.strip():
                metadata = {
                    'paragraph_count': len(doc.paragraphs),
                    'word_count': len(full_text.split()),
                    'char_count': len(full_text),
                    'extraction_method': 'python-docx'
                }
                
                return True, full_text, metadata
            
        except Exception as e1:
            self.logger.warning(f"python-docx failed: {e1}")
        
        # Fallback to mammoth
        try:
            with open(file_path, 'rb') as f:
                result = mammoth.extract_raw_text(f)
                full_text = result.value
                
                metadata = {
                    'word_count': len(full_text.split()),
                    'char_count': len(full_text),
                    'extraction_method': 'mammoth',
                    'conversion_messages': [str(msg) for msg in result.messages]
                }
                
                return True, full_text, metadata
                
        except Exception as e2:
                            return False, f"DOCX extraction failed: {str(e2)}", {}
    
    def analyze_figure(self, file_path: str) -> Tuple[bool, Dict[str, Any]]:
        """Analyze figure specifications and quality"""
        
        if not PIL_AVAILABLE:
            return False, {"error": "Image processing libraries not available"}
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._analyze_pdf_figure(file_path)
            elif file_ext in {'.png', '.jpg', '.jpeg', '.tiff', '.tif'}:
                return self._analyze_image_figure(file_path)
            else:
                return False, {"error": f"Unsupported figure type: {file_ext}"}
                
        except Exception as e:
            self.logger.error(f"Error analyzing figure {file_path}: {e}")
            return False, {"error": f"Analysis failed: {str(e)}"}
    
    def _analyze_image_figure(self, file_path: Path) -> Tuple[bool, Dict[str, Any]]:
        """Analyze image figure specifications"""
        
        with Image.open(file_path) as img:
            # Basic properties
            width, height = img.size
            mode = img.mode
            format_name = img.format
            
            # Calculate DPI
            dpi = img.info.get('dpi', (72, 72))
            if isinstance(dpi, tuple):
                dpi_x, dpi_y = dpi
            else:
                dpi_x = dpi_y = dpi
            
            # File size
            file_size = file_path.stat().st_size
            
            # Color analysis
            if img.mode == 'RGB':
                stat = ImageStat.Stat(img)
                is_grayscale = all(abs(stat.mean[i] - stat.mean[0]) < 10 for i in range(3))
            else:
                is_grayscale = img.mode in ['L', 'P', '1']
            
            # Quality assessment
            quality_issues = []
            
            if min(dpi_x, dpi_y) < 300:
                quality_issues.append(f"Low DPI: {min(dpi_x, dpi_y)} (recommended: 300+)")
            
            if width * height > 10000 * 10000:
                quality_issues.append(f"Very large dimensions: {width}x{height}")
            elif width * height < 600 * 400:
                quality_issues.append(f"Small dimensions: {width}x{height}")
            
            if file_size > 50 * 1024 * 1024:  # 50MB
                quality_issues.append(f"Large file size: {file_size / (1024*1024):.1f}MB")
            
            analysis = {
                'width': width,
                'height': height,
                'dpi_x': dpi_x,
                'dpi_y': dpi_y,
                'min_dpi': min(dpi_x, dpi_y),
                'color_mode': mode,
                'format': format_name,
                'file_size_mb': file_size / (1024 * 1024),
                'is_grayscale': is_grayscale,
                'aspect_ratio': round(width / height, 2),
                'megapixels': round((width * height) / 1000000, 1),
                'quality_issues': quality_issues,
                'meets_300_dpi': min(dpi_x, dpi_y) >= 300,
                'suitable_for_print': min(dpi_x, dpi_y) >= 300 and file_size < 50 * 1024 * 1024
            }
            
            return True, analysis
    
    def _analyze_pdf_figure(self, file_path: Path) -> Tuple[bool, Dict[str, Any]]:
        """Analyze PDF figure specifications"""
        
        if not PDF_AVAILABLE:
            return False, {"error": "PDF processing not available"}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    return False, {"error": "Empty PDF"}
                
                page = pdf.pages[0]
                
                # Get page dimensions (in points, 72 points = 1 inch)
                width_pt = float(page.width)
                height_pt = float(page.height)
                
                # Convert to inches and pixels
                width_in = width_pt / 72
                height_in = height_pt / 72
                width_px_300dpi = int(width_in * 300)
                height_px_300dpi = int(height_in * 300)
                
                file_size = file_path.stat().st_size
                
                quality_issues = []
                
                # Check dimensions
                if width_in < 3 or height_in < 3:
                    quality_issues.append(f"Small figure size: {width_in:.1f}x{height_in:.1f} inches")
                
                if width_in > 20 or height_in > 20:
                    quality_issues.append(f"Very large figure: {width_in:.1f}x{height_in:.1f} inches")
                
                if file_size > 10 * 1024 * 1024:  # 10MB for PDF
                    quality_issues.append(f"Large file size: {file_size / (1024*1024):.1f}MB")
                
                # Check for embedded images
                has_images = bool(page.images)
                
                analysis = {
                    'page_count': len(pdf.pages),
                    'width_points': width_pt,
                    'height_points': height_pt,
                    'width_inches': round(width_in, 2),
                    'height_inches': round(height_in, 2),
                    'width_px_at_300dpi': width_px_300dpi,
                    'height_px_at_300dpi': height_px_300dpi,
                    'file_size_mb': file_size / (1024 * 1024),
                    'has_embedded_images': has_images,
                    'is_vector': not has_images,
                    'aspect_ratio': round(width_pt / height_pt, 2),
                    'quality_issues': quality_issues,
                    'suitable_for_publication': len(quality_issues) == 0 and file_size < 10 * 1024 * 1024
                }
                
                return True, analysis
                
        except Exception as e:
            return False, {"error": f"PDF analysis failed: {str(e)}"}
    
    def get_figure_requirements_check(self, analysis: Dict[str, Any], target_journal: str = 'Nature') -> Dict[str, Any]:
        """Check figure against specific journal requirements"""
        
        # Journal requirements database
        requirements = {
            'Nature': {
                'min_dpi': 300,
                'preferred_formats': ['PDF', 'PNG', 'TIFF'],
                'max_file_size_mb': 30,
                'min_width_inches': 3.5,
                'max_width_inches': 7.0
            },
            'Science': {
                'min_dpi': 300,
                'preferred_formats': ['PDF', 'TIFF', 'EPS'],
                'max_file_size_mb': 20,
                'min_width_inches': 3.0,
                'max_width_inches': 6.5
            },
            'Cell': {
                'min_dpi': 300,
                'preferred_formats': ['PDF', 'TIFF'],
                'max_file_size_mb': 50,
                'min_width_inches': 3.5,
                'max_width_inches': 8.5
            },
            'The Lancet': {
                'min_dpi': 300,
                'preferred_formats': ['TIFF', 'PDF'],
                'max_file_size_mb': 25,
                'min_width_inches': 3.0,
                'max_width_inches': 6.0
            }
        }
        
        req = requirements.get(target_journal, requirements['Nature'])
        
        compliance = {
            'journal': target_journal,
            'meets_dpi': analysis.get('min_dpi', 0) >= req['min_dpi'],
            'meets_format': analysis.get('format', '').upper() in req['preferred_formats'],
            'meets_size_limit': analysis.get('file_size_mb', 0) <= req['max_file_size_mb'],
            'meets_width_range': req['min_width_inches'] <= analysis.get('width_inches', 0) <= req['max_width_inches'],
            'overall_compliant': True
        }
        
        # Check overall compliance
        compliance['overall_compliant'] = all([
            compliance['meets_dpi'],
            compliance['meets_format'], 
            compliance['meets_size_limit'],
            compliance['meets_width_range']
        ])
        
        return compliance
    
    def validate_research_paper(self, text: str) -> Dict[str, Any]:
        """Validate research paper structure and content"""
        
        validation = {
            'word_count': len(text.split()),
            'char_count': len(text),
            'has_abstract': False,
            'has_introduction': False,
            'has_methods': False,
            'has_results': False,
            'has_discussion': False,
            'has_references': False,
            'citation_count': 0,
            'section_headers': [],
            'issues': []
        }
        
        text_lower = text.lower()
        
        # Check for common sections
        if 'abstract' in text_lower:
            validation['has_abstract'] = True
        
        if 'introduction' in text_lower:
            validation['has_introduction'] = True
        
        if any(word in text_lower for word in ['method', 'methodology', 'materials']):
            validation['has_methods'] = True
        
        if 'result' in text_lower:
            validation['has_results'] = True
        
        if 'discussion' in text_lower or 'conclusion' in text_lower:
            validation['has_discussion'] = True
        
        if 'reference' in text_lower or 'bibliography' in text_lower:
            validation['has_references'] = True
        
        # Extract citations
        citations = self.extract_citations_from_text(text)
        validation['citation_count'] = len(citations)
        
        # Find section headers (simple heuristic)
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if len(line) < 50 and len(line) > 3:
                if any(keyword in line.lower() for keyword in ['abstract', 'introduction', 'method', 'result', 'discussion', 'conclusion', 'reference']):
                    validation['section_headers'].append(line)
        
        # Identify issues
        if validation['word_count'] < 1000:
            validation['issues'].append(f"Very short paper: {validation['word_count']} words")
        elif validation['word_count'] > 15000:
            validation['issues'].append(f"Very long paper: {validation['word_count']} words")
        
        if not validation['has_abstract']:
            validation['issues'].append("No abstract found")
        
        if validation['citation_count'] < 5:
            validation['issues'].append(f"Few citations: {validation['citation_count']}")
        
        required_sections = ['has_introduction', 'has_methods', 'has_results', 'has_discussion']
        missing_sections = [section.replace('has_', '') for section in required_sections if not validation[section]]
        if missing_sections:
            validation['issues'].append(f"Missing sections: {', '.join(missing_sections)}")
        
        return validation
    
    def extract_citations_from_text(self, text: str) -> List[Dict[str, str]]:
        """Extract citations from text using various patterns"""
        
        citations = []
        
        # Pattern 1: (Author, Year) format
        pattern1 = r'\(([^)]*(?:\d{4})[^)]*)\)'
        matches1 = re.findall(pattern1, text)
        for match in matches1:
            if any(char.isdigit() for char in match):
                citations.append({
                    'type': 'parenthetical',
                    'text': match,
                    'format': 'author_year'
                })
        
        # Pattern 2: [Number] format  
        pattern2 = r'\[(\d+(?:[-,]\d+)*)\]'
        matches2 = re.findall(pattern2, text)
        for match in matches2:
            citations.append({
                'type': 'numbered',
                'text': match,
                'format': 'numeric'
            })
        
        # Pattern 3: Author et al. format
        pattern3 = r'([A-Z][a-z]+ et al\.?,?\s*(?:\d{4})?)'
        matches3 = re.findall(pattern3, text)
        for match in matches3:
            citations.append({
                'type': 'narrative',
                'text': match,
                'format': 'author_year'
            })
        
        return citations
    
    def create_thumbnail(self, file_path: str, max_size: Tuple[int, int] = (200, 200)) -> Tuple[bool, str]:
        """Create thumbnail for image files"""
        
        if not PIL_AVAILABLE:
            return False, "Image processing not available"
        
        try:
            file_path = Path(file_path)
            thumbnail_path = self.upload_path / "thumbnails" / f"{file_path.stem}_thumb{file_path.suffix}"
            
            with Image.open(file_path) as img:
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                img.save(thumbnail_path, optimize=True, quality=85)
            
            return True, str(thumbnail_path)
            
        except Exception as e:
            self.logger.error(f"Thumbnail creation failed: {e}")
            return False, f"Thumbnail creation failed: {str(e)}"
    
    def convert_figure_format(self, file_path: str, target_format: str, 
                            target_dpi: int = 300) -> Tuple[bool, str, Dict[str, Any]]:
        """Convert figure to different format with specified DPI"""
        
        if not PIL_AVAILABLE:
            return False, "Image processing not available", {}
        
        try:
            file_path = Path(file_path)
            output_path = file_path.with_suffix(f'.{target_format.lower()}')
            
            with Image.open(file_path) as img:
                # Convert mode if necessary
                if target_format.upper() in ['JPEG', 'JPG'] and img.mode in ['RGBA', 'LA']:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                elif target_format.upper() == 'PNG' and img.mode not in ['RGB', 'RGBA']:
                    img = img.convert('RGBA')
                
                # Set DPI and save
                save_kwargs = {'dpi': (target_dpi, target_dpi)}
                
                if target_format.upper() == 'JPEG':
                    save_kwargs['quality'] = 95
                elif target_format.upper() == 'TIFF':
                    save_kwargs['compression'] = 'lzw'
                
                img.save(output_path, format=target_format.upper(), **save_kwargs)
            
            # Verify the conversion
            conversion_success, new_analysis = self.analyze_figure(str(output_path))
            
            return True, str(output_path), {
                'original_format': file_path.suffix,
                'new_format': target_format,
                'target_dpi': target_dpi,
                'actual_dpi': new_analysis.get('min_dpi') if conversion_success else None,
                'file_size_mb': output_path.stat().st_size / (1024 * 1024)
            }
            
        except Exception as e:
            self.logger.error(f"Format conversion failed: {e}")
            return False, f"Conversion failed: {str(e)}", {}
    
    def compress_file(self, file_path: str, quality: int = 85) -> Tuple[bool, str, Dict[str, Any]]:
        """Compress file to reduce size while maintaining quality"""
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        if file_ext in {'.jpg', '.jpeg', '.png'} and PIL_AVAILABLE:
            return self._compress_image(file_path, quality)
        elif file_ext == '.pdf':
            return self._compress_pdf(file_path)
        else:
            return False, f"Compression not supported for {file_ext}", {}
    
    def _compress_image(self, file_path: Path, quality: int) -> Tuple[bool, str, Dict[str, Any]]:
        """Compress image file"""
        try:
            output_path = file_path.parent / f"{file_path.stem}_compressed{file_path.suffix}"
            
            with Image.open(file_path) as img:
                # Convert RGBA to RGB for JPEG
                if img.mode in ['RGBA', 'LA'] and file_path.suffix.lower() in ['.jpg', '.jpeg']:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                save_kwargs = {'optimize': True}
                if file_path.suffix.lower() in ['.jpg', '.jpeg']:
                    save_kwargs['quality'] = quality
                
                img.save(output_path, **save_kwargs)
            
            original_size = file_path.stat().st_size
            compressed_size = output_path.stat().st_size
            compression_ratio = (1 - compressed_size / original_size) * 100
            
            return True, str(output_path), {
                'original_size_mb': original_size / (1024 * 1024),
                'compressed_size_mb': compressed_size / (1024 * 1024),
                'compression_ratio': round(compression_ratio, 1),
                'quality_setting': quality
            }
            
        except Exception as e:
            return False, f"Image compression failed: {str(e)}", {}
    
    def _compress_pdf(self, file_path: Path) -> Tuple[bool, str, Dict[str, Any]]:
        """Compress PDF file (basic implementation)"""
        # Note: Full PDF compression requires specialized libraries like PyPDF4 or qpdf
        return False, "PDF compression not implemented (requires additional libraries)", {}
    
    def check_file_integrity(self, file_path: str) -> Tuple[bool, str]:
        """Check if file is not corrupted and readable"""
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            # Basic existence and size check
            if not file_path.exists():
                return False, "File does not exist"
            
            if file_path.stat().st_size == 0:
                return False, "File is empty"
            
            # Format-specific integrity checks
            if file_ext == '.pdf':
                return self._check_pdf_integrity(file_path)
            elif file_ext in {'.png', '.jpg', '.jpeg', '.tiff', '.tif'}:
                return self._check_image_integrity(file_path)
            elif file_ext in {'.docx'}:
                return self._check_docx_integrity(file_path)
            elif file_ext in {'.txt'}:
                return self._check_text_integrity(file_path)
            else:
                return True, "File exists and has content"
                
        except Exception as e:
            return False, f"Integrity check failed: {str(e)}"
    
    def _check_pdf_integrity(self, file_path: Path) -> Tuple[bool, str]:
        """Check PDF file integrity"""
        if not PDF_AVAILABLE:
            return True, "PDF check skipped (libraries not available)"
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)
                
                if page_count == 0:
                    return False, "PDF has no pages"
                
                # Try to read first page
                first_page = pdf_reader.pages[0]
                first_page.extract_text()
                
                return True, f"PDF is valid ({page_count} pages)"
                
        except Exception as e:
            return False, f"PDF corruption detected: {str(e)}"
    
    def _check_image_integrity(self, file_path: Path) -> Tuple[bool, str]:
        """Check image file integrity"""
        if not PIL_AVAILABLE:
            return True, "Image check skipped (PIL not available)"
        
        try:
            with Image.open(file_path) as img:
                img.verify()
                
            # Re-open to get basic info
            with Image.open(file_path) as img:
                width, height = img.size
                mode = img.mode
                
                return True, f"Image is valid ({width}x{height}, {mode} mode)"
                
        except Exception as e:
            return False, f"Image corruption detected: {str(e)}"
    
    def _check_docx_integrity(self, file_path: Path) -> Tuple[bool, str]:
        """Check DOCX file integrity"""
        if not DOCX_AVAILABLE:
            return True, "DOCX check skipped (libraries not available)"
        
        try:
            doc = docx.Document(file_path)
            paragraph_count = len(doc.paragraphs)
            
            return True, f"DOCX is valid ({paragraph_count} paragraphs)"
            
        except Exception as e:
            return False, f"DOCX corruption detected: {str(e)}"
    
    def _check_text_integrity(self, file_path: Path) -> Tuple[bool, str]:
        """Check text file integrity"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1024)  # Read first 1KB
                
            if not content.strip():
                return False, "Text file appears to be empty"
            
            return True, "Text file is readable"
            
        except UnicodeDecodeError:
            return False, "Text file has encoding issues"
        except Exception as e:
            return False, f"Text file error: {str(e)}"
    
    def process_batch_upload(self, files_data: List[Tuple[bytes, str]], session_id: str, 
                           file_type: str = 'paper') -> List[Dict[str, Any]]:
        """Process multiple file uploads at once"""
        results = []
        
        for file_data, filename in files_data:
            success, file_path_or_error, file_info = self.save_uploaded_file(
                file_data, filename, session_id, file_type
            )
            
            result = {
                'filename': filename,
                'success': success,
                'message': file_path_or_error if not success else 'Upload successful',
                'file_info': file_info if success else {}
            }
            
            # For figures, also run analysis
            if success and file_type == 'figure':
                analysis_success, analysis_result = self.analyze_figure(file_path_or_error)
                result['analysis'] = {
                    'success': analysis_success,
                    'result': analysis_result
                }
            
            results.append(result)
        
        return results